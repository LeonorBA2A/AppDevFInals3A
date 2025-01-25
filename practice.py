import streamlit as st
import base64  # Import the base64 module
from docx import Document

# Path to your PDF file
pdf_file = r"C:\Users\Bien\Documents\Bien\School\3rd Year\Frist Sem\APDEV\ACM\SYSTEM\Py streamlit\Minimal Professional Resume.pdf"

# Read the PDF as binary and embed it in an iframe
with open(pdf_file, "rb") as file:
    pdf_data = file.read()
    base64_pdf = base64.b64encode(pdf_data).decode('utf-8')

pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
st.markdown(pdf_display, unsafe_allow_html=True)

# Load the Word document
doc = Document(r"C:\Users\Bien\Documents\Bien\School\3rd Year\Frist Sem\APDEV\ACM\SYSTEM\Py streamlit\Adhi Gopalam - SM.docx")

# Display the content
for paragraph in doc.paragraphs:
    st.write(paragraph.text)

uploaded_file = st.file_uploader("Upload a document", type=["pdf", "txt", "docx"])

if uploaded_file is not None:
    if uploaded_file.name.endswith(".pdf"):
        # Handle PDF display
        base64_pdf = base64.b64encode(uploaded_file.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)
    elif uploaded_file.name.endswith(".txt"):
        # Handle TXT display
        content = uploaded_file.read().decode("utf-8")
        st.text(content)
    elif uploaded_file.name.endswith(".docx"):
        # Handle DOCX display
        from docx import Document
        doc = Document(uploaded_file)
        for paragraph in doc.paragraphs:
            st.write(paragraph.text)
